import os
import csv
import re
import sys
import threading
import time
from pathlib import Path
import webview
from docx import Document
import win32com.client

# Word SaveAs PDF format constant
WD_FORMAT_PDF = 17

def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

def start_backend_server():
    """Start the FastAPI backend server in a background thread"""
    import uvicorn
    
    # Ensure _MEIPASS is on sys.path so 'import backend' works inside PyInstaller
    try:
        meipass = sys._MEIPASS
        if meipass not in sys.path:
            sys.path.insert(0, meipass)
    except AttributeError:
        pass
    
    try:
        from backend import app as fastapi_app
        from starlette.responses import FileResponse, Response
        
        # Serve the built frontend static files via a catch-all route
        # (mounted AFTER API routes so /api/* is handled first)
        static_dir = resource_path(os.path.join('frontend', 'build'))
        if os.path.isdir(static_dir):
            @fastapi_app.get("/{full_path:path}")
            async def serve_spa(full_path: str):
                # Try to serve the exact file
                file_path = os.path.join(static_dir, full_path)
                if os.path.isfile(file_path):
                    # Determine content type
                    import mimetypes
                    content_type, _ = mimetypes.guess_type(file_path)
                    return FileResponse(file_path, media_type=content_type)
                # Fallback to index.html for SPA routing
                index_path = os.path.join(static_dir, "index.html")
                if os.path.isfile(index_path):
                    return FileResponse(index_path, media_type="text/html")
                return Response(status_code=404)
        
        config = uvicorn.Config(fastapi_app, host="127.0.0.1", port=8000, log_level="warning")
        server = uvicorn.Server(config)
        server.run()
    except Exception as e:
        import traceback
        log_path = os.path.join(os.path.expanduser('~'), 'DocumentForge_Error.log')
        with open(log_path, 'w') as f:
            f.write("Failed to start backend server:\n")
            f.write(traceback.format_exc())

class Api:
    def __init__(self):
        self._window = None

    def set_window(self, window):
        self._window = window

    def select_file(self, file_type):
        """Invoke native file picker"""
        if file_type == 'csv':
            file_types = ('CSV Files (*.csv)', 'All files (*.*)')
        elif file_type == 'docx':
            file_types = ('Word Documents (*.docx)', 'All files (*.*)')
        else:
            file_types = ('All files (*.*)',)

        result = self._window.create_file_dialog(webview.OPEN_DIALOG, allow_multiple=False, file_types=file_types)
        if result:
            return result[0]
        return None
        
    def select_folder(self):
         """Invoke native folder picker"""
         result = self._window.create_file_dialog(webview.FOLDER_DIALOG, allow_multiple=False)
         if result:
             return result[0]
         return None

    def get_metadata(self, csv_path, docx_path):
        """Parse CSV headers and DOCX placeholders"""
        try:
            # 1. Parse CSV Headers
            print(f"Reading CSV: {csv_path}")
            headers = []
            rows_count = 0
            with open(csv_path, "r", encoding="utf-8") as f:
                reader = csv.reader(f)
                headers_raw = next(reader, [])
                # Apply the same normalization as the original script
                headers = [re.sub(r"\s+", "", h.strip()) for h in headers_raw]
                headers.append("P_ ADDRESS") # Support the edge case from original script
                
                rows_count = sum(1 for row in reader if any(v.strip() for v in row))
                
            # 2. Parse DOCX Placeholders
            print(f"Reading DOCX: {docx_path}")
            doc = Document(docx_path)
            placeholders = set()
            
            # Helper to extract #WORDS
            def extract_from_text(text):
                # Find words starting with #
                found = re.findall(r'#\w+(?:\s+\w+)?', text)
                for f in found:
                     placeholders.add(f)
            
            for paragraph in doc.paragraphs:
                extract_from_text(paragraph.text)
                
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            extract_from_text(paragraph.text)
                            
            # Also check headers/footers
            for section in doc.sections:
                for header_footer in (section.header, section.footer):
                    if header_footer is not None:
                        for paragraph in header_footer.paragraphs:
                            extract_from_text(paragraph.text)

            return {
                "success": True,
                "csv_headers": sorted(list(set(headers))),
                "docx_placeholders": sorted(list(placeholders)),
                "total_rows": rows_count
            }
        except Exception as e:
            print(f"Error parsing metadata: {e}")
            return {"success": False, "error": str(e)}

    # Generation Engine
    def generate(self, csv_path, docx_path, output_dir, mapping, output_format):
        """Run generation in background to not freeze UI"""
        thread = threading.Thread(
            target=self._run_generation, 
            args=(csv_path, docx_path, output_dir, mapping, output_format)
        )
        thread.start()
        return True
        
    def _run_generation(self, csv_path, docx_path, output_dir, mapping, output_format):
        try:
            output_path = Path(output_dir)
            output_path.mkdir(parents=True, exist_ok=True)
            
            # Use original logic
            records = self._read_csv(Path(csv_path))
            total = len(records)
            
            # Map placeholders
            self._window.evaluate_js(f"window.updateProgress(0, {total}, 'Starting generation...')")
            
            # If PDF is requested, open Word once
            word_app = None
            if output_format in ['pdf', 'both']:
                self._window.evaluate_js(f"window.updateProgress(0, {total}, 'Starting Microsoft Word...')")
                try: # pypiwin32 doesn't work if Word isn't installed
                    word_app = win32com.client.Dispatch("Word.Application")
                    word_app.Visible = False
                    word_app.DisplayAlerts = False
                except Exception as e:
                    self._window.evaluate_js(f"window.showError('Could not launch Word for PDF conversion: {e}')")
                    return

            for i, record in enumerate(records, start=1):
                # Build replacements based on user mapping
                replacements = {}
                for docx_tag, csv_col in mapping.items():
                    # Handle the P_ ADDRESS edgecase seamlessly if mapped
                    lookup_col = "P_ADDRESS" if csv_col == "P_ ADDRESS"  else csv_col 
                    val = record.get(lookup_col, "")
                    replacements[docx_tag] = val
                    
                # Sort longest key first
                replacements = dict(sorted(replacements.items(), key=lambda kv: len(kv[0]), reverse=True))

                name = record.get("NAME", f"row_{i}").strip().replace(" ", "_")
                if not name: name = f"row_{i}"
                filename_base = f"{i:04d}_{name}"
                
                # DOCX Process
                try:
                    out_docx = output_path / f"{filename_base}.docx"
                    doc = self._fill_template(Path(docx_path), replacements)
                    doc.save(out_docx)
                    
                    # PDF Process
                    if word_app is not None and output_format in ['pdf', 'both']:
                        self._window.evaluate_js(f"window.updateProgress({i}, {total}, 'Converting {filename_base} to PDF...')")
                        out_pdf = output_path / f"{filename_base}.pdf"
                        
                        pdf_doc = word_app.Documents.Open(str(out_docx.absolute()))
                        pdf_doc.SaveAs(str(out_pdf.absolute()), FileFormat=WD_FORMAT_PDF)
                        pdf_doc.Close(False)
                        
                        # Cleanup DOCX if they only wanted PDF
                        if output_format == 'pdf':
                            os.remove(out_docx)
                            
                    else:
                         self._window.evaluate_js(f"window.updateProgress({i}, {total}, 'Generated {filename_base}.docx')")

                except Exception as e:
                     print(f"Error on row {i}: {e}")
                     
            if word_app:
                word_app.Quit()

            self._window.evaluate_js(f"window.generationComplete('{output_dir}')")
            
        except Exception as e:
            self._window.evaluate_js(f"window.showError('{str(e)}')")
            print(f"Fatal error: {e}")

    # Adapted from original script
    def _read_csv(self, path: Path):
        rows = []
        with open(path, "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                if all(v.strip() == "" for v in row.values()):
                    continue
                record = {}
                for header, value in row.items():
                    key = re.sub(r"\s+", "", header.strip())
                    record[key] = value.strip() if value else ""
                rows.append(record)
        return rows
        
    def _replace_in_paragraph(self, paragraph, replacements):
        full_text = paragraph.text
        if "#" not in full_text: return

        for run in paragraph.runs:
            for placeholder, value in replacements.items():
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)

        remaining = paragraph.text
        for placeholder, value in replacements.items():
            if placeholder in remaining:
                new_text = remaining
                for ph, val in replacements.items():
                    new_text = new_text.replace(ph, val)
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                    for run in paragraph.runs[1:]:
                        run.text = ""
                break

    def _replace_in_table(self, table, replacements):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._replace_in_paragraph(paragraph, replacements)

    def _fill_template(self, template_path, replacements):
        doc = Document(template_path)
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)
        for table in doc.tables:
            self._replace_in_table(table, replacements)
        for section in doc.sections:
            for header_footer in (section.header, section.footer):
                if header_footer is not None:
                    for paragraph in header_footer.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)
        return doc


if __name__ == '__main__':
    api = Api()
    
    is_dev = len(sys.argv) > 1 and sys.argv[1] == '--dev'
    
    if is_dev:
        # Development mode — use Vite dev server
        html_path = 'http://localhost:5173'
    else:
        # Production mode — start FastAPI backend serving static files
        backend_thread = threading.Thread(target=start_backend_server, daemon=True)
        backend_thread.start()
        
        # Wait for server to be ready
        import socket
        for _ in range(50):  # 5 seconds max
            try:
                with socket.create_connection(("127.0.0.1", 8000), timeout=0.1):
                    break
            except (ConnectionRefusedError, socket.timeout):
                time.sleep(0.1)
        
        html_path = 'http://127.0.0.1:8000'
    
    print(f"Loading UI from: {html_path}")
    
    window = webview.create_window(
        'Document Forge', 
        url=html_path,
        js_api=api,
        width=1200, 
        height=850,
        min_size=(900, 600),
        text_select=False,
        background_color='#020617'
    )
    api.set_window(window)
    webview.start(debug=is_dev)

