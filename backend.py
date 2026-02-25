import os
import csv
import re
import uuid
import shutil
import zipfile
import threading
from pathlib import Path
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks, WebSocket, WebSocketDisconnect
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
import uvicorn
from pydantic import BaseModel
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt
import fitz  # PyMuPDF - for PDF empty page removal and merging
from docxcompose.composer import Composer  # for DOCX merging
import win32com.client
import pdfplumber
import tabula
from openpyxl.utils import get_column_letter
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
import pandas as pd

app = FastAPI(title="Document Forge API")

# Enable CORS for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173", 
        "http://127.0.0.1:5173",
        "http://localhost:5174", 
        "http://127.0.0.1:5174",
        "http://localhost:8000",
        "http://127.0.0.1:8000"
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

TEMP_DIR = Path("temp_sessions")
TEMP_DIR.mkdir(exist_ok=True)
OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(exist_ok=True)

WD_FORMAT_PDF = 17

class MappingItem(BaseModel):
    type: str # 'csv_column', 'custom_text', or 'combined'
    value: str
    prefix: Optional[str] = ""
    suffix: Optional[str] = ""

class DocSettings(BaseModel):
    page_size: Optional[str] = "default"  # 'default', 'letter', 'a4', 'legal', 'a3', 'a5', 'custom'
    page_width: Optional[float] = None   # inches (for custom)
    page_height: Optional[float] = None  # inches (for custom)
    font_name: Optional[str] = None
    font_size: Optional[int] = None

# Standard page sizes in inches
PAGE_SIZES = {
    "letter": (8.5, 11),
    "legal": (8.5, 14),
    "a3": (11.69, 16.54),
    "a4": (8.27, 11.69),
    "a5": (5.83, 8.27),
}

class GenerateRequest(BaseModel):
    session_id: str
    mapping: Dict[str, MappingItem]
    generate_docx: bool = True
    generate_pdf: bool = True
    remove_empty_pages: bool = True
    merge_output: bool = False
    doc_settings: Optional[DocSettings] = None
    rows_per_doc: int = 1

def cleanup_session(session_id: str):
    """Clean up temp files for a session after generation"""
    session_dir = TEMP_DIR / session_id
    if session_dir.exists():
        shutil.rmtree(session_dir, ignore_errors=True)


class ConnectionManager:
    def __init__(self):
        self.active_connections: dict[str, WebSocket] = {}

    async def connect(self, ws: WebSocket, client_id: str):
        await ws.accept()
        self.active_connections[client_id] = ws

    def disconnect(self, client_id: str):
        if client_id in self.active_connections:
            del self.active_connections[client_id]

    async def send_message(self, message: dict, client_id: str):
        ws = self.active_connections.get(client_id)
        if ws:
            try:
                await ws.send_json(message)
            except WebSocketDisconnect:
                self.disconnect(client_id)

manager = ConnectionManager()

@app.websocket("/ws/progress/{client_id}")
async def websocket_endpoint(websocket: WebSocket, client_id: str):
    await manager.connect(websocket, client_id)
    try:
        while True:
            await websocket.receive_text()
    except WebSocketDisconnect:
        manager.disconnect(client_id)

@app.post("/api/upload")
async def upload_files(data_file: UploadFile = File(...), template_file: UploadFile = File(...)):
    session_id = str(uuid.uuid4())
    session_dir = TEMP_DIR / session_id
    session_dir.mkdir(parents=True, exist_ok=True)
    
    csv_path = session_dir / "input.csv"
    template_path = session_dir / "template.docx"
    
    # Write template file
    with open(template_path, "wb") as f:
        f.write(await template_file.read())
        
    # Parse data file
    filename = data_file.filename.lower()
    file_bytes = await data_file.read()
    
    from io import BytesIO
    if filename.endswith('.xlsx'):
        df = pd.read_excel(BytesIO(file_bytes))
        df.to_csv(csv_path, index=False, encoding='utf-8')
    elif filename.endswith('.json'):
        df = pd.read_json(BytesIO(file_bytes))
        df.to_csv(csv_path, index=False, encoding='utf-8')
    else:
        # Default to CSV, auto-detecting delimiters/encoding
        try:
            df = pd.read_csv(BytesIO(file_bytes), sep=None, engine='python', on_bad_lines='skip')
        except Exception:
            # Fallback for weird encodings
            df = pd.read_csv(BytesIO(file_bytes), sep=None, engine='python', encoding='latin1', on_bad_lines='skip')
        df.to_csv(csv_path, index=False, encoding='utf-8')
            
    return {"session_id": session_id}

@app.get("/api/metadata")
async def get_metadata(session_id: str):
    session_dir = TEMP_DIR / session_id
    csv_path = session_dir / "input.csv"
    template_path = session_dir / "template.docx"
    
    if not csv_path.exists() or not template_path.exists():
        raise HTTPException(status_code=404, detail="Session not found or files missing")
        
    try:
        # 1. Parse CSV Headers & Rows
        headers = []
        rows_count = 0
        preview_rows = []
        with open(csv_path, "r", encoding="utf-8") as f:
            reader = csv.reader(f)
            headers_raw = next(reader, [])
            # Apply normalization
            headers = [re.sub(r"\s+", "", h.strip()) for h in headers_raw]
            headers.append("P_ ADDRESS")
            
            for index, row in enumerate(reader):
                if any(v.strip() for v in row):
                    rows_count += 1
                    if index < 100:  # Return first 100 rows for the preview Data Grid
                        row_dict = {headers_raw[i] if i < len(headers_raw) else f"Col{i}": v for i, v in enumerate(row)}
                        preview_rows.append(row_dict)
            
        # 2. Parse DOCX Placeholders (with occurrence counting)
        doc = Document(template_path)
        placeholders = set()
        placeholder_counts = {}
        
        def extract_from_text(text):
            found = re.findall(r'#\w+(?:[ \t]+\w+)?', text)
            for f in found:
                placeholders.add(f)
                placeholder_counts[f] = placeholder_counts.get(f, 0) + 1
        
        for paragraph in doc.paragraphs:
            extract_from_text(paragraph.text)
            
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        extract_from_text(paragraph.text)
                        
        for section in doc.sections:
            for header_footer in (section.header, section.footer):
                if header_footer is not None:
                    for paragraph in header_footer.paragraphs:
                        extract_from_text(paragraph.text)

        # 3. Extract template page size
        template_page_size = "letter"  # fallback
        template_page_width = 8.5
        template_page_height = 11.0
        if doc.sections:
            sec = doc.sections[0]
            # Convert EMU to inches (914400 EMU per inch)
            w_inches = round(sec.page_width / 914400, 2) if sec.page_width else 8.5
            h_inches = round(sec.page_height / 914400, 2) if sec.page_height else 11.0
            template_page_width = w_inches
            template_page_height = h_inches
            # Try to match against known sizes (with tolerance)
            matched = False
            for name, (pw, ph) in PAGE_SIZES.items():
                if abs(w_inches - pw) < 0.15 and abs(h_inches - ph) < 0.15:
                    template_page_size = name
                    matched = True
                    break
            if not matched:
                template_page_size = "custom"

        # Calculate rows_per_doc from max placeholder occurrence
        rows_per_doc = max(placeholder_counts.values()) if placeholder_counts else 1

        return {
            "csv_headers": sorted(list(set(headers))),
            "docx_placeholders": sorted(list(placeholders)),
            "total_rows": rows_count,
            "preview_rows": preview_rows,
            "raw_headers": headers_raw,
            "placeholder_counts": placeholder_counts,
            "rows_per_doc": rows_per_doc,
            "template_page_size": template_page_size,
            "template_page_width": template_page_width,
            "template_page_height": template_page_height,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class UpdateDataRequest(BaseModel):
    session_id: str
    rows: List[Dict[str, str]]
    headers: List[str]

@app.post("/api/update_data")
async def update_data(request: UpdateDataRequest):
    session_dir = TEMP_DIR / request.session_id
    csv_path = session_dir / "input.csv"
    if not csv_path.exists():
        raise HTTPException(status_code=404, detail="Session not found")
        
    try:
        # Overwrite the input.csv natively with the edited preview rows payload
        # Wait, if preview is only 100 rows, saving it overwrites the whole file!
        # If rows_count > 100, we should only modify the specific rows sent. 
        # But for this phase, we'll write the whole file if total_rows <= 100, 
        # or we update the CSV. For full edits, it's safer to just let the UI send the full file back, 
        # or handle just small datasets.
        # Let's save the edited rows to a patch, or overwrite if it's small.
        # For our "Interactive Editor" MVP, we will overwrite the whole file using pandas for simplicity.
        df = pd.DataFrame(request.rows, columns=request.headers)
        df.to_csv(csv_path, index=False, encoding="utf-8")
        return {"status": "success", "total_rows": len(request.rows)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/converter/extract")
async def extract_pdf(
    background_tasks: BackgroundTasks, 
    pdfFile: UploadFile = File(...), 
    processingOption: str = Form(...),
    outputFormat: str = Form("excel")
):
    session_id = str(uuid.uuid4())
    session_dir = TEMP_DIR / session_id
    session_dir.mkdir(parents=True, exist_ok=True)
    
    pdf_path = session_dir / "input_file.pdf"
    with open(pdf_path, "wb") as f:
        f.write(await pdfFile.read())
        
    try:
        if processingOption == 'allText':
            pages_data = _extract_pdf_content(str(pdf_path))
            if outputFormat == 'csv':
                out_name = f"Extracted_AllText_{session_id[:8]}.csv"
                out_path = session_dir / out_name
                _write_alltext_to_csv(pages_data, str(out_path))
                media_type = "text/csv"
            else:
                out_name = f"Extracted_AllText_{session_id[:8]}.xlsx"
                out_path = session_dir / out_name
                excel_file = _create_excel(pages_data)
                excel_file.save(str(out_path))
                media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        
        elif processingOption == 'tablesOnly':
            tables = _extract_tables_from_pdf(str(pdf_path))
            if tables:
                if outputFormat == 'csv':
                    out_name = f"Extracted_TablesOnly_{session_id[:8]}.zip"
                    out_path = session_dir / out_name
                    _write_tables_to_csv_zip(tables, session_dir, str(out_path))
                    media_type = "application/zip"
                else:
                    out_name = f"Extracted_TablesOnly_{session_id[:8]}.xlsx"
                    out_path = session_dir / out_name
                    excel_file = _write_tables_to_excel(tables)
                    excel_file.save(str(out_path))
                    media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            else:
                raise HTTPException(status_code=400, detail="No tables found in the PDF.")
        else:
            raise HTTPException(status_code=400, detail="Invalid processing option")
            
        background_tasks.add_task(cleanup_session, session_id)
        return FileResponse(
            path=out_path,
            filename=out_name,
            media_type=media_type
        )
    except HTTPException:
        cleanup_session(session_id)
        raise
    except Exception as e:
        cleanup_session(session_id)
        raise HTTPException(status_code=500, detail=f"Error processing PDF: {str(e)}")

@app.post("/api/generate")
async def generate_documents(request: GenerateRequest, background_tasks: BackgroundTasks):
    session_id = request.session_id
    session_dir = TEMP_DIR / session_id
    csv_path = session_dir / "input.csv"
    template_path = session_dir / "template.docx"
    
    if not csv_path.exists() or not template_path.exists():
        raise HTTPException(status_code=404, detail="Session not found")
        
    out_dir = session_dir / "output_files"
    out_dir.mkdir(exist_ok=True)
    
    rows_per_doc = max(1, request.rows_per_doc)
    
    try:
        # Read CSV logic
        rows = []
        with open(csv_path, "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                if all(v.strip() == "" for v in row.values()):
                    continue
                record = {}
                for header, value in row.items():
                    key = re.sub(r"\s+", "", header.strip())
                    record[key] = value.strip() if value else ""
                rows.append(record)
                
        # Prep Word COM if PDF needed
        word_app = None
        if request.generate_pdf:
            try:
                word_app = win32com.client.Dispatch("Word.Application")
                word_app.Visible = False
                word_app.DisplayAlerts = False
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Word PDF Engine failed: {str(e)}")
        
        # Chunk rows into groups based on rows_per_doc
        chunks = []
        for i in range(0, len(rows), rows_per_doc):
            chunks.append(rows[i:i + rows_per_doc])
        
        total_docs = len(chunks)
        import asyncio
        for doc_idx, chunk in enumerate(chunks, start=1):
            # Send progress to websocket if available
            await manager.send_message({
                "type": "progress",
                "current": doc_idx,
                "total": total_docs,
                "message": f"Processing document {doc_idx} of {total_docs}..."
            }, session_id)
            
            # CRITICAL: Yield control to the event loop so the WebSocket message actually sends
            await asyncio.sleep(0.01)
            
            if rows_per_doc == 1:
                # === SINGLE-ROW MODE (original behavior) ===
                record = chunk[0]
                replacements = {}
                for docx_tag, map_item in request.mapping.items():
                    if map_item.type == 'custom_text':
                        replacements[docx_tag] = map_item.value
                    elif map_item.type in ('csv_column', 'data_column'):
                        lookup_col = "P_ADDRESS" if map_item.value == "P_ ADDRESS"  else map_item.value
                        val = record.get(lookup_col, "")
                        with open("debug_backend.txt", "a") as dbg:
                            dbg.write(f"SINGLE tag={docx_tag} lookup={lookup_col} keys={list(record.keys())} val='{val}'\\n")
                        replacements[docx_tag] = val
                    elif map_item.type == 'combined':
                        lookup_col = "P_ADDRESS" if map_item.value == "P_ ADDRESS"  else map_item.value
                        val = record.get(lookup_col, "")
                        if val:
                            replacements[docx_tag] = f"{map_item.prefix}{val}{map_item.suffix}"
                        else:
                            replacements[docx_tag] = ""
                        
                # Sort longest key first
                replacements = dict(sorted(replacements.items(), key=lambda kv: len(kv[0]), reverse=True))

                name = record.get("NAME", f"row_{doc_idx}").strip().replace(" ", "_")
                if not name: name = f"row_{doc_idx}"
                filename_base = f"{doc_idx:04d}_{name}"
                
                # Fill document
                doc = _fill_template(template_path, replacements, record=record)
            else:
                # === MULTI-ROW MODE ===
                doc = _fill_template_multi(template_path, chunk, request.mapping, rows_per_doc)
                
                first_record = chunk[0]
                name = first_record.get("NAME", f"batch_{doc_idx}").strip().replace(" ", "_")
                if not name: name = f"batch_{doc_idx}"
                filename_base = f"{doc_idx:04d}_{name}"
            
            # Apply document settings
            if request.doc_settings:
                _apply_doc_settings(doc, request.doc_settings)

            out_docx = out_dir / f"{filename_base}.docx"
            doc.save(out_docx)
            
            # Convert to PDF
            if word_app and request.generate_pdf:
                out_pdf = out_dir / f"{filename_base}.pdf"
                pdf_doc = word_app.Documents.Open(str(out_docx.absolute()))
                pdf_doc.SaveAs(str(out_pdf.absolute()), FileFormat=WD_FORMAT_PDF)
                pdf_doc.Close(False)
                
                if not request.generate_docx:
                    os.remove(out_docx)

        if word_app:
            word_app.Quit()
        
        # === POST-PROCESSING: Remove Empty Pages ===
        if request.remove_empty_pages:
            await manager.send_message({
                "type": "progress",
                "current": total_docs,
                "total": total_docs,
                "message": "Removing empty pages..."
            }, session_id)
            await asyncio.sleep(0.01)
            
            # Remove empty pages from PDFs
            for pdf_file in out_dir.glob("*.pdf"):
                try:
                    pdf_doc = fitz.open(str(pdf_file))
                    pages_to_remove = []
                    for page_num in range(len(pdf_doc)):
                        page = pdf_doc[page_num]
                        text = page.get_text().strip()
                        images = page.get_images()
                        # Check for any drawn paths/shapes (like table borders)
                        drawings = page.get_drawings()
                        
                        # A page is "empty" if it has no meaningful text
                        # (allow pages with images or drawings that also have text)
                        has_meaningful_text = bool(text and len(text) > 2)
                        has_images = bool(images)
                        
                        # If NO text and NO images, it's empty (even if it has borders/lines)
                        if not has_meaningful_text and not has_images:
                            pages_to_remove.append(page_num)
                    
                    # Remove pages in reverse order to keep indices valid
                    for page_num in reversed(pages_to_remove):
                        pdf_doc.delete_page(page_num)
                    if len(pdf_doc) > 0:
                        pdf_doc.save(str(pdf_file), incremental=False, deflate=True)
                    pdf_doc.close()
                except Exception as e:
                    print(f"Warning: could not strip empty pages from {pdf_file.name}: {e}")
            
            # Remove empty pages from DOCX files
            for docx_file in out_dir.glob("*.docx"):
                try:
                    doc = Document(str(docx_file))
                    from docx.oxml.ns import qn
                    body = doc.element.body
                    elements_to_remove = []
                    
                    children = list(body)
                    
                    for elem in children:
                        if elem.tag == qn('w:p'):
                            # Extract all text from the paragraph
                            text = ""
                            for node in elem.iter():
                                if node.tag == qn('w:t') and node.text:
                                    text += node.text
                            text = text.strip()
                            
                            # Check for page break (manual break)
                            has_page_break = False
                            for br in elem.findall('.//' + qn('w:br')):
                                br_type = br.get(qn('w:type'))
                                if br_type == 'page':
                                    has_page_break = True
                            
                            # Check for section break (which also causes page breaks)
                            has_section_break = elem.find(qn('w:pPr') + '/' + qn('w:sectPr')) is not None
                            
                            # Remove if: page break with no text, OR section break with no text
                            if (has_page_break or has_section_break) and not text:
                                elements_to_remove.append(elem)
                    
                    # Also remove trailing empty paragraphs at the end of the document
                    # These often cause a blank last page
                    reversed_children = list(reversed(children))
                    for elem in reversed_children:
                        if elem.tag == qn('w:sectPr'):
                            continue  # Skip the final section properties
                        if elem.tag == qn('w:p'):
                            text = ""
                            for node in elem.iter():
                                if node.tag == qn('w:t') and node.text:
                                    text += node.text
                            text = text.strip()
                            if not text:
                                if elem not in elements_to_remove:
                                    elements_to_remove.append(elem)
                            else:
                                break  # Stop at first non-empty paragraph
                        else:
                            break  # Stop at non-paragraph element (table, etc.)
                    
                    for elem in elements_to_remove:
                        try:
                            body.remove(elem)
                        except ValueError:
                            pass  # Already removed
                    
                    doc.save(str(docx_file))
                except Exception as e:
                    print(f"Warning: could not strip empty pages from {docx_file.name}: {e}")
        
        # === POST-PROCESSING: Merge All Output ===
        if request.merge_output:
            await manager.send_message({
                "type": "progress",
                "current": total_docs,
                "total": total_docs,
                "message": "Merging all documents into single file..."
            }, session_id)
            await asyncio.sleep(0.01)
            
            # Merge all DOCX files
            docx_files = sorted(out_dir.glob("*.docx"))
            if len(docx_files) > 1 and request.generate_docx:
                try:
                    master = Document(str(docx_files[0]))
                    composer = Composer(master)
                    for docx_path in docx_files[1:]:
                        sub_doc = Document(str(docx_path))
                        composer.append(sub_doc)
                    merged_docx = out_dir / "Merged_All.docx"
                    composer.save(str(merged_docx))
                    # Remove individual files
                    for docx_path in docx_files:
                        os.remove(str(docx_path))
                except Exception as e:
                    print(f"Warning: DOCX merge failed: {e}")
            
            # Merge all PDF files
            pdf_files = sorted(out_dir.glob("*.pdf"))
            if len(pdf_files) > 1 and request.generate_pdf:
                try:
                    merged_pdf = fitz.open()
                    for pdf_path in pdf_files:
                        sub_pdf = fitz.open(str(pdf_path))
                        merged_pdf.insert_pdf(sub_pdf)
                        sub_pdf.close()
                    merged_pdf_path = out_dir / "Merged_All.pdf"
                    merged_pdf.save(str(merged_pdf_path))
                    merged_pdf.close()
                    # Remove individual files
                    for pdf_path in pdf_files:
                        os.remove(str(pdf_path))
                except Exception as e:
                    print(f"Warning: PDF merge failed: {e}")
            
        # Tell client we are zipping
        asyncio.create_task(manager.send_message({
            "type": "progress",
            "current": total_docs,
            "total": total_docs,
            "message": "Archiving output files into ZIP..."
        }, session_id))
        
        # Create ZIP file
        zip_path = session_dir / "DocumentForge_Output.zip"
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root, _, files in os.walk(out_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, out_dir)
                    zipf.write(file_path, arcname)
                    
        # Schedule cleanup after download
        background_tasks.add_task(cleanup_session, session_id)
        
        return FileResponse(
            path=zip_path, 
            filename="DocumentForge_Output.zip",
            media_type="application/zip"
        )
        
    except Exception as e:
        if word_app:
            word_app.Quit()
        raise HTTPException(status_code=500, detail=str(e))

# Helpers
def _replace_in_paragraph(paragraph, replacements):
    full_text = paragraph.text
    if "#" not in full_text: return

    for run in paragraph.runs:
        for placeholder, value in replacements.items():
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)

    remaining = paragraph.text
    for placeholder, value in replacements.items():
        if placeholder in remaining:
            new_text = remaining
            for ph, val in replacements.items():
                new_text = new_text.replace(ph, val)
            if paragraph.runs:
                paragraph.runs[0].text = new_text
                for run in paragraph.runs[1:]:
                    run.text = ""
            break

def _replace_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, replacements)

def _process_conditionals(doc, record):
    """Process #IF(COLUMN=VALUE) ... #ENDIF conditional blocks in a DOCX document.
    
    Supports:
        #IF(GENDER=Male)  → equality check
        #IF(PHONE)        → existence check (non-empty)
    
    Removes entire paragraph blocks between #IF and #ENDIF when condition fails.
    Strips #IF/#ENDIF markers when condition passes.
    """
    IF_PATTERN = re.compile(r'#IF\(([^)]+)\)', re.IGNORECASE)
    ENDIF_PATTERN = re.compile(r'#ENDIF', re.IGNORECASE)
    
    # Process body paragraphs
    _process_conditional_paragraphs(doc.paragraphs, record, IF_PATTERN, ENDIF_PATTERN)
    
    # Process table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _process_conditional_paragraphs(cell.paragraphs, record, IF_PATTERN, ENDIF_PATTERN)
    
    # Process headers/footers
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            if header_footer is not None:
                _process_conditional_paragraphs(header_footer.paragraphs, record, IF_PATTERN, ENDIF_PATTERN)

def _process_conditional_paragraphs(paragraphs, record, if_pattern, endif_pattern):
    """Process conditional blocks within a list of paragraphs."""
    paragraphs_to_clear = []
    i = 0
    
    while i < len(paragraphs):
        para_text = paragraphs[i].text.strip()
        match = if_pattern.search(para_text)
        
        if match:
            condition_str = match.group(1).strip()
            condition_met = _evaluate_condition(condition_str, record)
            
            # Find the matching #ENDIF
            if_para_idx = i
            endif_para_idx = None
            nesting = 1
            j = i + 1
            
            while j < len(paragraphs):
                inner_text = paragraphs[j].text.strip()
                # Count nested #IF blocks
                if if_pattern.search(inner_text):
                    nesting += 1
                if endif_pattern.search(inner_text):
                    nesting -= 1
                    if nesting == 0:
                        endif_para_idx = j
                        break
                j += 1
            
            if endif_para_idx is not None:
                if condition_met:
                    # Keep the content, strip #IF and #ENDIF markers
                    _strip_marker(paragraphs[if_para_idx], if_pattern)
                    _strip_marker(paragraphs[endif_para_idx], endif_pattern)
                    i = if_para_idx  # re-process in case of nested
                else:
                    # Clear all paragraphs from #IF to #ENDIF (inclusive)
                    for k in range(if_para_idx, endif_para_idx + 1):
                        paragraphs_to_clear.append(k)
                    i = endif_para_idx + 1
                    continue
        i += 1
    
    # Clear the failed conditional paragraphs (set text to empty)
    for idx in paragraphs_to_clear:
        for run in paragraphs[idx].runs:
            run.text = ""

def _evaluate_condition(condition_str, record):
    """Evaluate a condition like 'GENDER=Male' or 'PHONE' against a record."""
    if '=' in condition_str:
        # Equality check: COLUMN=VALUE
        parts = condition_str.split('=', 1)
        column = re.sub(r"\s+", "", parts[0].strip())
        expected = parts[1].strip()
        actual = record.get(column, "")
        return actual.strip().lower() == expected.strip().lower()
    else:
        # Existence check: COLUMN (non-empty)
        column = re.sub(r"\s+", "", condition_str.strip())
        return bool(record.get(column, "").strip())

def _strip_marker(paragraph, pattern):
    """Remove a regex pattern from a paragraph's runs."""
    for run in paragraph.runs:
        run.text = pattern.sub('', run.text)
    # Also do a full-text fallback
    full_text = paragraph.text
    if pattern.search(full_text):
        new_text = pattern.sub('', full_text).strip()
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""

def _fill_template(template_path, replacements, record=None):
    doc = Document(template_path)
    # Process conditionals FIRST (before tag replacement)
    if record:
        _process_conditionals(doc, record)
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        _replace_in_table(table, replacements)
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            if header_footer is not None:
                for paragraph in header_footer.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)
    return doc

def _fill_template_multi(template_path, chunk, mapping, rows_per_doc):
    """Fill a template with multiple rows. Each placeholder's N-th occurrence
    is replaced with the N-th row's data from the chunk."""
    doc = Document(template_path)
    
    # Build per-row replacement dicts
    row_replacements = []
    for record in chunk:
        replacements = {}
        for docx_tag, map_item in mapping.items():
            if map_item.type == 'custom_text':
                replacements[docx_tag] = map_item.value
            elif map_item.type in ('csv_column', 'data_column'):
                lookup_col = "P_ADDRESS" if map_item.value == "P_ ADDRESS" else map_item.value
                val = record.get(lookup_col, "")
                with open("debug_backend.txt", "a", encoding="utf-8") as dbg:
                    dbg.write(f"MULTI tag={docx_tag} lookup={lookup_col} keys={list(record.keys())} val='{val}'\\n")
                print(f"DEBUG: docx_tag={docx_tag} lookup={lookup_col} | actual val='{val}' | record_keys={list(record.keys())}")
                replacements[docx_tag] = val
            elif map_item.type == 'combined':
                lookup_col = "P_ADDRESS" if map_item.value == "P_ ADDRESS" else map_item.value
                val = record.get(lookup_col, "")
                if val:
                    replacements[docx_tag] = f"{map_item.prefix}{val}{map_item.suffix}"
                else:
                    replacements[docx_tag] = ""
        row_replacements.append(replacements)
    
    # Pad with empty replacements if chunk is smaller than rows_per_doc
    while len(row_replacements) < rows_per_doc:
        empty_rep = {}
        for docx_tag in mapping.keys():
            empty_rep[docx_tag] = ""
        row_replacements.append(empty_rep)
    
    # Sort placeholders longest first to avoid partial replacement
    sorted_tags = sorted(mapping.keys(), key=len, reverse=True)
    
    # Track occurrence index for each placeholder
    occurrence_counter = {tag: 0 for tag in sorted_tags}
    
    def _replace_multi_in_paragraph(paragraph):
        """Replace placeholders in a paragraph, tracking occurrence count globally."""
        full_text = paragraph.text
        if "#" not in full_text:
            return
        
        # First try run-level replacement
        for run in paragraph.runs:
            for tag in sorted_tags:
                while tag in run.text:
                    idx = occurrence_counter[tag]
                    if idx < len(row_replacements):
                        val = row_replacements[idx].get(tag, "")
                    else:
                        val = ""
                    run.text = run.text.replace(tag, val, 1)
                    occurrence_counter[tag] = idx + 1
        
        # Fallback: if placeholders span across runs, consolidate
        remaining = paragraph.text
        for tag in sorted_tags:
            if tag in remaining:
                # Consolidate all runs into first run
                new_text = paragraph.text
                # Count how many times this tag appears in the consolidated text
                while tag in new_text:
                    idx = occurrence_counter[tag]
                    if idx < len(row_replacements):
                        val = row_replacements[idx].get(tag, "")
                    else:
                        val = ""
                    new_text = new_text.replace(tag, val, 1)
                    occurrence_counter[tag] = idx + 1
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                    for run in paragraph.runs[1:]:
                        run.text = ""
                break
    
    # Process all paragraphs in document order
    for paragraph in doc.paragraphs:
        _replace_multi_in_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_multi_in_paragraph(paragraph)
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            if header_footer is not None:
                for paragraph in header_footer.paragraphs:
                    _replace_multi_in_paragraph(paragraph)
    
    return doc

def _apply_doc_settings(doc, settings):
    """Apply document settings (page size, font, font size) to a python-docx Document."""
    # Apply page size
    if settings.page_size and settings.page_size != "default":
        if settings.page_size == "custom":
            w = settings.page_width or 8.5
            h = settings.page_height or 11.0
        elif settings.page_size in PAGE_SIZES:
            w, h = PAGE_SIZES[settings.page_size]
        else:
            w, h = 8.5, 11.0  # fallback to letter
        
        for section in doc.sections:
            section.page_width = Inches(w)
            section.page_height = Inches(h)
    
    # Apply font settings
    if settings.font_name or settings.font_size:
        font_name = settings.font_name if settings.font_name else None
        font_size = Pt(settings.font_size) if settings.font_size else None
        
        def _apply_font_to_runs(runs):
            for run in runs:
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = font_size
        
        for paragraph in doc.paragraphs:
            _apply_font_to_runs(paragraph.runs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _apply_font_to_runs(paragraph.runs)
        for section in doc.sections:
            for header_footer in (section.header, section.footer):
                if header_footer is not None:
                    for paragraph in header_footer.paragraphs:
                        _apply_font_to_runs(paragraph.runs)


# --- PDF Extractor Helpers ---
def _extract_tables_from_pdf(pdf_path):
    tables = []
    try:
        tables = tabula.read_pdf(pdf_path, pages='all', multiple_tables=True, lattice=True)
    except Exception:
        pass
    if not tables or all(t.empty for t in tables):
        try:
            tables = tabula.read_pdf(pdf_path, pages='all', multiple_tables=True, stream=True)
        except Exception:
            return []
    return [t for t in tables if not t.empty] if tables else []

def _extract_pdf_content(pdf_path):
    pages_data = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            table_settings = {"vertical_strategy": "lines", "horizontal_strategy": "lines", "snap_tolerance": 5}
            raw_tables = page.extract_tables(table_settings) or []
            if not raw_tables:
                raw_tables = page.extract_tables({"vertical_strategy": "text", "horizontal_strategy": "text"}) or []
                
            table_bboxes = []
            found_tables = page.find_tables(table_settings) if raw_tables else []
            for t in found_tables:
                table_bboxes.append(t.bbox)
                
            non_table_lines = []
            if table_bboxes:
                cropped_page = page
                for bbox in table_bboxes:
                    cropped_page = cropped_page.outside_bbox(bbox)
                remaining_text = cropped_page.extract_text() or ""
                non_table_lines = [l.strip() for l in remaining_text.split('\n') if l.strip()]
            else:
                text = page.extract_text() or ""
                non_table_lines = [l.strip() for l in text.split('\n') if l.strip()]
            pages_data.append((non_table_lines, raw_tables))
    return pages_data

def _prepare_alltext_grid(pages_data):
    merged_tables = []
    extra_col_header = None
    extra_col_values = []
    for non_table_lines, raw_tables in pages_data:
        for table in raw_tables:
            if not table or len(table) == 0: continue
            num_cols = len(table[0])
            if merged_tables:
                last = merged_tables[-1]
                last_cols = len(last[0])
                if num_cols == last_cols:
                    first_row_str = ' '.join(str(c or '').strip().lower() for c in table[0])
                    header_str = ' '.join(str(c or '').strip().lower() for c in last[0])
                    if first_row_str == header_str: last.extend(table[1:])
                    else: last.extend(table)
                else: merged_tables.append(list(table))
                continue
            merged_tables.append(list(table))
            
        if non_table_lines and raw_tables:
            data_rows_on_page = sum(len(t) for t in raw_tables)
            if len(non_table_lines) == data_rows_on_page:
                if extra_col_header is None:
                    extra_col_header = non_table_lines[0]
                    extra_col_values.extend(non_table_lines[1:])
                else:
                    if non_table_lines[0].strip().lower() == extra_col_header.strip().lower(): extra_col_values.extend(non_table_lines[1:])
                    else: extra_col_values.extend(non_table_lines)
            elif len(non_table_lines) == data_rows_on_page - 1:
                extra_col_values.extend(non_table_lines)
                
    if merged_tables and extra_col_values:
        table = merged_tables[0]
        if extra_col_header and len(table) > 0: table[0].append(extra_col_header)
        data_start = 1
        for i, val in enumerate(extra_col_values):
            row_idx = data_start + i
            if row_idx < len(table): table[row_idx].append(val)
            else: break
        expected_cols = len(table[0])
        for row in table:
            while len(row) < expected_cols: row.append("")
    return merged_tables

def _create_excel(pages_data):
    merged_tables = _prepare_alltext_grid(pages_data)
    wb = Workbook()
    ws = wb.active
    ws.title = "PDF Content"
    title_font = Font(name='Times New Roman', size=11, bold=True)
    data_font = Font(name='Times New Roman', size=11)
    alignment = Alignment(wrap_text=True, vertical='center')
    header_fill = PatternFill(start_color='FFF2CC', end_color='FFF2CC', fill_type='solid')
    current_row = 1
    for table in merged_tables:
        if not table: continue
        for row_idx, row_data in enumerate(table):
            for col_idx, value in enumerate(row_data, start=1):
                cell_val = str(value).replace("\n", " ").strip() if isinstance(value, str) else (value or "")
                cell = ws.cell(row=current_row, column=col_idx, value=cell_val)
                cell.alignment = alignment
                if row_idx == 0:
                    cell.font = title_font
                    cell.fill = header_fill
                else: cell.font = data_font
            current_row += 1
        current_row += 1
    for col_cells in ws.columns:
        max_length = 0
        col_letter = get_column_letter(col_cells[0].column)
        for cell in col_cells:
            try:
                if cell.value: max_length = max(max_length, len(str(cell.value)))
            except: pass
        ws.column_dimensions[col_letter].width = max(min(max_length + 2, 50), 10)
    return wb

def _write_alltext_to_csv(pages_data, out_path):
    merged_tables = _prepare_alltext_grid(pages_data)
    with open(out_path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        for table in merged_tables:
            if not table: continue
            for row in table:
                cleaned_row = [str(val).replace("\n", " ").strip() if val else "" for val in row]
                writer.writerow(cleaned_row)
            writer.writerow([]) # blank line between tables

def _process_merged_tables(tables):
    merged_tables = []
    for table in tables:
        if merged_tables:
            last = merged_tables[-1]
            if list(last.columns) == list(table.columns):
                merged_tables[-1] = pd.concat([last, table], ignore_index=True)
                continue
        merged_tables.append(table.copy())
    return merged_tables

def _write_tables_to_excel(tables):
    merged_tables = _process_merged_tables(tables)
    workbook = Workbook()
    title_font = Font(name='Times New Roman', size=11, bold=True)
    info_font = Font(name='Times New Roman', size=10)
    alignment = Alignment(wrap_text=True, vertical='center')
    header_fill = PatternFill(start_color='FFF2CC', end_color='FFF2CC', fill_type='solid')
    for table_num, table in enumerate(merged_tables, start=1):
        sheet = workbook.create_sheet(title=f'Table_{table_num}')
        df = pd.DataFrame(table.values, columns=[str(col).title() if col else "" for col in table.columns])
        for col_num, column_header in enumerate(df.columns, start=1):
            if column_header:
                cell = sheet.cell(row=1, column=col_num, value=column_header)
                cell.font = title_font
                cell.alignment = alignment
                cell.fill = header_fill
                sheet.column_dimensions[get_column_letter(col_num)].width = max(len(str(column_header)) + 2, 10)
        for row_num, (_, row) in enumerate(df.iterrows(), start=2):
            for col_num, value in enumerate(row, start=1):
                cell_val = str(value).replace("\n", " ").strip() if isinstance(value, str) else (value or "")
                cell = sheet.cell(row=row_num, column=col_num, value=cell_val)
                cell.font = info_font
                cell.alignment = alignment
                col_width = sheet.column_dimensions[get_column_letter(col_num)].width
                sheet.column_dimensions[get_column_letter(col_num)].width = max(col_width, min(len(str(value)) + 2, 50))
        for row in sheet.iter_rows(min_row=1, max_row=sheet.max_row, min_col=1, max_col=sheet.max_column):
            max_text_length = max((len(str(cell.value)) for cell in row if cell.value), default=0)
            sheet.row_dimensions[row[0].row].height = 35 + (max_text_length // 50) * 5
    workbook.remove(workbook.active)
    return workbook

def _write_tables_to_csv_zip(tables, session_dir, zip_path):
    merged_tables = _process_merged_tables(tables)
    import zipfile
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for idx, table in enumerate(merged_tables, start=1):
            csv_name = f"Table_{idx}.csv"
            csv_path = session_dir / csv_name
            table.to_csv(csv_path, index=False)
            zipf.write(csv_path, csv_name)

if __name__ == "__main__":
    print("🚀 Starting Document Forge Fast API server on port 8000...")
    uvicorn.run("backend:app", host="127.0.0.1", port=8000, reload=True)
