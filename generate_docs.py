"""
CSV → DOCX Template Filler
============================
Reads rows from a CSV file and generates one filled DOCX per row by
replacing #COLUMN_NAME placeholders in the template with actual values.

Usage:
    python generate_docs.py

Folder layout expected:
    sheets/input.csv               – data source (first row = headers)
    doc_templete/PLA Petition new.docx  – Word template with #PLACEHOLDERS
    output/                        – generated docs land here (auto-created)
"""

import csv
import re
from pathlib import Path

from docx import Document


# ── Configuration ────────────────────────────────────────────────────────────
BASE_DIR      = Path(__file__).resolve().parent
CSV_PATH      = BASE_DIR / "sheets" / "input.csv"
TEMPLATE_PATH = BASE_DIR / "doc_templete" / "PLA Petition new.docx"
OUTPUT_DIR    = BASE_DIR / "output"


def read_csv(path: Path) -> list[dict[str, str]]:
    """Return a list of dicts, one per data row. Keys are column headers."""
    rows = []
    with open(path, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            # Skip completely empty rows
            if all(v.strip() == "" for v in row.values()):
                continue
            # Normalise headers: strip spaces, collapse whitespace
            record = {}
            for header, value in row.items():
                key = re.sub(r"\s+", "", header.strip())
                record[key] = value.strip() if value else ""
            rows.append(record)
    return rows


def replace_in_paragraph(paragraph, replacements: dict[str, str]):
    """Replace #PLACEHOLDER tokens in a paragraph, preserving formatting."""
    full_text = paragraph.text
    # Quick check – skip paragraphs with no placeholders
    if "#" not in full_text:
        return

    # Strategy: try run-level replacement first (preserves formatting).
    # If a placeholder spans multiple runs, fall back to full-paragraph rebuild.
    for run in paragraph.runs:
        for placeholder, value in replacements.items():
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)

    # Verify all placeholders were replaced; if not, do a full-text replace
    remaining = paragraph.text
    for placeholder, value in replacements.items():
        if placeholder in remaining:
            # Full-text fallback: merge everything into the first run
            new_text = remaining
            for ph, val in replacements.items():
                new_text = new_text.replace(ph, val)
            if paragraph.runs:
                paragraph.runs[0].text = new_text
                for run in paragraph.runs[1:]:
                    run.text = ""
            break


def replace_in_table(table, replacements: dict[str, str]):
    """Replace placeholders in every cell of a table."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, replacements)


def fill_template(template_path: Path, replacements: dict[str, str]) -> Document:
    """Open a fresh copy of the template and replace all placeholders."""
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        replace_in_table(table, replacements)

    # Also handle headers / footers
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            if header_footer is not None:
                for paragraph in header_footer.paragraphs:
                    replace_in_paragraph(paragraph, replacements)

    return doc


def build_replacements(record: dict[str, str]) -> dict[str, str]:
    """Build a {#KEY: value} mapping from a data record.
    Sorted longest-key-first so #TOSW is replaced before #TOS, etc."""
    raw = {f"#{key}": value for key, value in record.items()}
    # Special case for P_ADDRESS as the template has '#P_ ADDRESS' with a space
    if "P_ADDRESS" in record:
        raw["#P_ ADDRESS"] = record["P_ADDRESS"]
    return dict(sorted(raw.items(), key=lambda kv: len(kv[0]), reverse=True))


def main():
    # Read data
    print(f"📖  Reading data from: {CSV_PATH}")
    records = read_csv(CSV_PATH)
    print(f"    Found {len(records)} data rows.\n")

    if not records:
        print("⚠️  No data rows found. Exiting.")
        return

    # Ensure output directory exists
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Show which placeholders will be used (from first row)
    sample = build_replacements(records[0])
    print("🔑  Placeholders to replace:")
    for ph in sorted(sample):
        print(f"    {ph}  →  (e.g. {sample[ph][:40]})")
    print()

    # Generate one document per row
    for i, record in enumerate(records, start=1):
        replacements = build_replacements(record)

        # Use NAME + serial for a human-friendly filename
        name = record.get("NAME", f"row_{i}").strip().replace(" ", "_")
        filename = f"{i:04d}_{name}.docx"
        output_path = OUTPUT_DIR / filename

        doc = fill_template(TEMPLATE_PATH, replacements)
        doc.save(output_path)

        print(f"  ✅  [{i}/{len(records)}]  {filename}")

    print(f"\n🎉  Done! {len(records)} documents saved to: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
